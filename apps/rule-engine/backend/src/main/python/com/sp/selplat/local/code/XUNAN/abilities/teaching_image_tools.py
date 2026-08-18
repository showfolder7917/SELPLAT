#!/usr/bin/env python3
"""中文教学 DOCX 解析、主题图片批量生成和清单导出能力。"""

from __future__ import annotations

import argparse
from dataclasses import asdict, dataclass
import hashlib
import json
from pathlib import Path
import re

from docx import Document
from PIL import Image, ImageDraw, ImageFont


INVALID_FILE_CHARACTERS = re.compile(r"[<>:\"/\\|?*\x00-\x1f]")


@dataclass(frozen=True)
class AnnotatedToken:
    pinyin: str
    text: str


@dataclass(frozen=True)
class AnnotatedLine:
    tokens: tuple[AnnotatedToken, ...]

    @property
    def text(self) -> str:
        return "".join(item.text for item in self.tokens)


@dataclass(frozen=True)
class TopicPoem:
    title: str
    dynasty: str
    author: str
    lines: tuple[AnnotatedLine, ...]


@dataclass(frozen=True)
class TopicContent:
    story: tuple[str, ...]
    interpretation: tuple[str, ...]
    core_ideas: tuple[str, ...]
    life_tips: tuple[str, ...]
    footer: str
    illustration: str
    study_illustration: str


def sanitize_file_name(value: str) -> str:
    """替换 Windows、macOS 和 Linux 都不稳定的文件名字符。"""

    return INVALID_FILE_CHARACTERS.sub("_", value).strip(" .") or "untitled"


def base_name(path: Path) -> str:
    """只移除最后一个扩展名，保留业务名称中的点。"""

    return path.stem


def parse_annotated_docx(source_path: Path) -> list[TopicPoem]:
    """解析“第一段拼音、第二段汉字”的单行表格并组装诗词。"""

    document = Document(source_path)
    lines: list[AnnotatedLine] = []
    for table in document.tables:
        if not table.rows:
            continue
        tokens: list[AnnotatedToken] = []
        for cell in table.rows[0].cells:
            paragraphs = cell.paragraphs
            pinyin = paragraphs[0].text.strip() if paragraphs else ""
            text = paragraphs[1].text if len(paragraphs) > 1 else paragraphs[0].text
            tokens.append(AnnotatedToken(pinyin, text))
        if any(token.text for token in tokens):
            lines.append(AnnotatedLine(tuple(tokens)))
    poems: list[TopicPoem] = []
    index = 0
    while index < len(lines):
        title_line = lines[index]
        if index + 1 >= len(lines):
            break
        attribution_line = lines[index + 1]
        match = re.fullmatch(r"【([^】]+)】(.+)", attribution_line.text)
        if not match:
            index += 1
            continue
        body: list[AnnotatedLine] = []
        index += 2
        while index < len(lines) and not re.fullmatch(r"【([^】]+)】(.+)", lines[index].text):
            if index + 1 < len(lines) and re.fullmatch(r"【([^】]+)】(.+)", lines[index + 1].text):
                break
            body.append(lines[index])
            index += 1
        poems.append(TopicPoem(title_line.text.replace(" ", ""), match.group(1), match.group(2), tuple(body)))
    return poems


def load_topic_content(path: Path) -> dict[str, TopicContent]:
    """读取审校 JSON；任何栏目缺失都保留为空并由调用方决定失败策略。"""

    raw = json.loads(path.read_text(encoding="utf-8"))
    topics = raw.get("topics") or {}
    result: dict[str, TopicContent] = {}
    for title, value in topics.items():
        result[title] = TopicContent(
            tuple(value.get("story") or ()),
            tuple(value.get("interpretation") or ()),
            tuple(value.get("coreIdeas") or ()),
            tuple(value.get("lifeTips") or ()),
            str(value.get("footer") or ""),
            str(value.get("illustration") or ""),
            str(value.get("studyIllustration") or value.get("illustration") or ""),
        )
    return result


def resolve_asset(content_path: Path, relative_path: str) -> Path:
    """只允许主题 JSON 同目录内的真实素材，阻断路径逃逸。"""

    root = content_path.parent.resolve()
    target = (root / relative_path).resolve()
    if target != root and root not in target.parents:
        raise ValueError(f"插画路径越界：{relative_path}")
    if not target.is_file():
        raise FileNotFoundError(f"主题插画不存在：{target}")
    return target


def _font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """优先选择系统中文字体，缺失时退回 Pillow 默认字体并保持程序可运行。"""

    candidates = [
        Path("/System/Library/Fonts/PingFang.ttc"),
        Path("/System/Library/Fonts/STHeiti Light.ttc"),
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    ]
    for candidate in candidates:
        if candidate.is_file():
            try:
                return ImageFont.truetype(str(candidate), size=size, index=1 if bold else 0)
            except OSError:
                continue
    return ImageFont.load_default()


def display_font_status() -> str:
    """返回明确字体状态，便于统一测试区分缓存缺失和渲染失败。"""

    font = _font(24)
    return "字体资源已加载" if isinstance(font, ImageFont.FreeTypeFont) else "字体资源不存在，使用系统回退字体"


def render_topic_image(
    poem: TopicPoem,
    content: TopicContent,
    illustration_path: Path,
    target: Path,
    *,
    overwrite: bool = False,
) -> None:
    """生成 16:9 教学主题图，插画和文字区域保持稳定分栏。"""

    if target.exists() and not overwrite:
        raise FileExistsError(f"目标已存在：{target}")
    canvas = Image.new("RGB", (1600, 900), "#F7F0E2")
    with Image.open(illustration_path) as source:
        image = source.convert("RGB")
        image.thumbnail((900, 900))
        canvas.paste(image, (700 + (900 - image.width) // 2, (900 - image.height) // 2))
    overlay = Image.new("RGBA", canvas.size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(overlay)
    draw.rounded_rectangle((60, 55, 760, 845), radius=42, fill=(255, 253, 247, 226))
    draw.text((110, 92), poem.title, font=_font(64, bold=True), fill="#5B3A22")
    draw.text((112, 175), f"【{poem.dynasty}】{poem.author}", font=_font(28), fill="#6E7776")
    y = 245
    for line in poem.lines:
        pinyin = " ".join(token.pinyin for token in line.tokens if token.pinyin)
        draw.text((112, y), pinyin, font=_font(20), fill="#75878B")
        draw.text((112, y + 34), line.text, font=_font(38), fill="#273E47")
        y += 92
    if content.footer:
        draw.text((112, 785), content.footer, font=_font(22), fill="#9B6653")
    canvas = Image.alpha_composite(canvas.convert("RGBA"), overlay).convert("RGB")
    target.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(target, "JPEG", quality=92, optimize=True)


def generate_batch(
    source_docx: Path,
    output_root: Path,
    topic_json: Path,
    *,
    limit: int = 0,
    overwrite: bool = False,
    missing_policy: str = "fail",
) -> dict[str, object]:
    """批量生成主题图；缺内容或缺插画时默认立即失败。"""

    poems = parse_annotated_docx(source_docx)
    contents = load_topic_content(topic_json)
    selected = poems[:limit] if limit > 0 else poems
    generated: list[dict[str, object]] = []
    for poem in selected:
        content = contents.get(poem.title)
        if content is None:
            if missing_policy == "skip":
                continue
            raise ValueError(f"缺少审校故事内容：{poem.title}")
        illustration = resolve_asset(topic_json, content.illustration)
        target = output_root / f"{sanitize_file_name(poem.title)}.jpg"
        render_topic_image(poem, content, illustration, target, overwrite=overwrite)
        generated.append({
            "title": poem.title,
            "output": str(target),
            "sha256": hashlib.sha256(target.read_bytes()).hexdigest(),
        })
    manifest_path = output_root / "生成清单.json"
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(generated, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {"source_count": len(selected), "image_count": len(generated), "output_root": str(output_root)}


def export_manifest(source_docx: Path, target_json: Path) -> dict[str, object]:
    """把核定版 DOCX 转成可供图片与 PPT 生成器复用的 JSON 清单。"""

    poems = parse_annotated_docx(source_docx)
    payload = {"poems": [asdict(poem) for poem in poems]}
    target_json.parent.mkdir(parents=True, exist_ok=True)
    target_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return {"poem_count": len(poems), "target": str(target_json)}


def main() -> int:
    """提供 batch 与 manifest 两个稳定命令。"""

    parser = argparse.ArgumentParser()
    subparsers = parser.add_subparsers(dest="action", required=True)
    batch = subparsers.add_parser("batch")
    batch.add_argument("source_docx", type=Path)
    batch.add_argument("output_root", type=Path)
    batch.add_argument("topic_json", type=Path)
    batch.add_argument("--limit", type=int, default=0)
    batch.add_argument("--overwrite", action="store_true")
    batch.add_argument("--missing-policy", choices=("fail", "skip"), default="fail")
    manifest = subparsers.add_parser("manifest")
    manifest.add_argument("source_docx", type=Path)
    manifest.add_argument("target_json", type=Path)
    arguments = parser.parse_args()
    if arguments.action == "batch":
        result = generate_batch(
            arguments.source_docx,
            arguments.output_root,
            arguments.topic_json,
            limit=arguments.limit,
            overwrite=arguments.overwrite,
            missing_policy=arguments.missing_policy,
        )
    else:
        result = export_manifest(arguments.source_docx, arguments.target_json)
    print(json.dumps({"status": "completed", **result}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
