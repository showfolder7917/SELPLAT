#!/usr/bin/env python3
"""中文拼音、古诗文解析与注音 DOCX 生成能力。

替代原 Java 拼音生成文件组，保留最长词组纠音、原文保护、朗读变调、
诗文元数据解析、逐字注音表格和源文件不可覆盖约束。
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
import hashlib
import json
from pathlib import Path
import re
import unicodedata

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

try:
    from pypinyin import Style, lazy_pinyin
except ImportError:  # 运行入口会给出明确依赖错误，模块仍可用于只依赖人工词典的测试。
    Style = None
    lazy_pinyin = None


CJK_PATTERN = re.compile(r"[\u3400-\u9fff]")
ATTRIBUTION_PATTERN = re.compile(r"(?:【([^】]+)】\s*([^★]+)|([^【】]+)【([^】]+)】)$")
VOLUME_PATTERN = re.compile(r"^[一二三四五六七八九]年级[上下]册[：:]?$")
DIRECTORY_PATTERN = re.compile(r"目录$")


@dataclass(frozen=True)
class PinyinCell:
    """保存一个原文字元和对应拼音；标点拼音为空。"""

    pinyin: str
    text: str


@dataclass(frozen=True)
class Poem:
    """保存一首诗的标准标题、朝代、作者和正文行。"""

    title: str
    dynasty: str
    author: str
    lines: tuple[str, ...]

    @property
    def attribution(self) -> str:
        return f"【{self.dynasty}】{self.author}" if self.dynasty else self.author


@dataclass(frozen=True)
class Article:
    """保存一篇文言文及其可视化拆行规则。"""

    title: str
    dynasty: str
    author: str
    paragraphs: tuple[str, ...]

    def display_lines(self, maximum_columns: int) -> list[str]:
        """按码点拆行，同时避免把“一、不”留在行尾破坏变调。"""

        result: list[str] = []
        for paragraph in self.paragraphs:
            remaining = paragraph
            while len(remaining) > maximum_columns:
                cut = maximum_columns
                if remaining[cut - 1] in {"一", "不"}:
                    cut -= 1
                result.append(remaining[:cut])
                remaining = remaining[cut:]
            if remaining:
                result.append(remaining)
        return result


def load_overrides(dictionary_path: Path | None) -> dict[str, tuple[str, ...]]:
    """读取 UTF-8 TSV 纠音词典，并阻断词长和音节数不一致。"""

    if dictionary_path is None:
        return {}
    overrides: dict[str, tuple[str, ...]] = {}
    for line_number, raw_line in enumerate(dictionary_path.read_text(encoding="utf-8").splitlines(), 1):
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        parts = line.split("\t")
        if len(parts) != 2:
            raise ValueError(f"纠音词典第 {line_number} 行必须是词语和拼音两列。")
        phrase = parts[0].strip()
        syllables = tuple(parts[1].split())
        if len(phrase) != len(syllables):
            raise ValueError(f"纠音词典第 {line_number} 行音节数不匹配：{phrase}")
        overrides[phrase] = syllables
    return overrides


class PinyinConverter:
    """把原文转换为逐字拼音单元，人工长词优先于通用读音。"""

    def __init__(self, overrides: dict[str, tuple[str, ...]] | None = None):
        self.overrides = dict(overrides or {})
        self._phrases = sorted(self.overrides, key=len, reverse=True)

    def convert(self, text: str) -> list[PinyinCell]:
        """转换完整字符串，标点、空格和原文字元顺序保持不变。"""

        normalized = unicodedata.normalize("NFC", str(text))
        cells: list[PinyinCell] = []
        overridden_indexes: set[int] = set()
        index = 0
        while index < len(normalized):
            phrase = next((item for item in self._phrases if normalized.startswith(item, index)), None)
            if phrase:
                cells.extend(PinyinCell(pinyin, character) for pinyin, character in zip(self.overrides[phrase], phrase))
                overridden_indexes.update(range(index, index + len(phrase)))
                index += len(phrase)
                continue
            character = normalized[index]
            if CJK_PATTERN.fullmatch(character):
                cells.append(PinyinCell(self._single_pinyin(character), character))
            else:
                cells.append(PinyinCell("", character))
            index += 1
        return self._apply_reading_sandhi(cells, overridden_indexes)

    @staticmethod
    def _single_pinyin(character: str) -> str:
        """调用标准 Python 拼音库；缺少依赖时禁止静默产生错误读音。"""

        if lazy_pinyin is None or Style is None:
            raise RuntimeError("缺少 pypinyin；请使用项目 Python 运行时安装 requirements-python.txt。")
        return lazy_pinyin(character, style=Style.TONE, neutral_tone_with_five=False, errors="default")[0]

    def _apply_reading_sandhi(
        self, cells: list[PinyinCell], overridden_indexes: set[int]
    ) -> list[PinyinCell]:
        """在未被人工词组覆盖的位置应用“一、不”朗读变调。"""

        result = list(cells)
        for index, cell in enumerate(cells):
            if cell.text not in {"一", "不"}:
                continue
            if index in overridden_indexes:
                continue
            following = next((item for item in cells[index + 1:] if item.pinyin), None)
            if not following:
                continue
            fourth_tone = following.pinyin.endswith(("à", "è", "ì", "ò", "ù", "ǜ"))
            if cell.text == "一":
                replacement = "yí" if fourth_tone else "yì"
            else:
                replacement = "bú" if fourth_tone else "bù"
            result[index] = PinyinCell(replacement, cell.text)
        return result


def _document_lines(source_path: Path) -> list[str]:
    """按文档顺序读取段落和表格文字。"""

    document = Document(source_path)
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            lines.append("".join(cell.text for cell in row.cells).strip())
    return lines


def parse_poetry(source_path: Path) -> list[Poem]:
    """从教材 DOCX 中识别标题、署名和正文，过滤目录与分册容器。"""

    lines = _document_lines(source_path)
    poems: list[Poem] = []
    title = ""
    dynasty = ""
    author = ""
    body: list[str] = []

    def flush() -> None:
        nonlocal title, dynasty, author, body
        if title and body:
            inferred_dynasty, inferred_author = _infer_known_attribution(title, dynasty, author)
            poems.append(Poem(_clean_title(title), inferred_dynasty, inferred_author, tuple(body)))
        title, dynasty, author, body = "", "", "", []

    body_started = False
    for raw_line in lines:
        line = raw_line.strip()
        if not line or line.isdigit() or DIRECTORY_PATTERN.search(line):
            continue
        if VOLUME_PATTERN.fullmatch(line):
            body_started = True
            continue
        if not body_started and "年级" in line:
            continue
        parsed = _parse_attribution(line)
        if parsed and title:
            dynasty, author = parsed
            continue
        if _looks_like_body(line):
            if title:
                body.append(line)
            continue
        if "【" in line and "】" in line:
            before = line.split("【", 1)[0].strip()
            parsed = _parse_attribution(line[len(before):].strip())
            if before and parsed:
                flush()
                title = before
                dynasty, author = parsed
                continue
        flush()
        title = line
    flush()
    return poems


def parse_classical_chinese(source_path: Path) -> list[Article]:
    """解析文言文篇目；已知无署名篇目使用稳定教材元数据。"""

    poems = parse_poetry(source_path)
    return [Article(item.title, item.dynasty, item.author, item.lines) for item in poems]


def _parse_attribution(line: str) -> tuple[str, str] | None:
    clean = re.sub(r"★.*$", "", line).strip()
    if clean == "汉乐府":
        return "汉", "乐府"
    match = ATTRIBUTION_PATTERN.fullmatch(clean)
    if not match:
        return None
    if match.group(1):
        return match.group(1).strip(), match.group(2).replace(" ", "").strip()
    return match.group(4).strip(), match.group(3).replace(" ", "").strip()


def _infer_known_attribution(title: str, dynasty: str, author: str) -> tuple[str, str]:
    known = {
        "月夜": ("唐", "刘方平"),
        "《弟子规·谨》（节选）": ("清", "李毓秀"),
    }
    return (dynasty, author) if dynasty or author else known.get(_clean_title(title), ("", "佚名"))


def _clean_title(title: str) -> str:
    return re.sub(r"★.*$", "", title).replace(" ", "").strip("：:")


def _looks_like_body(line: str) -> bool:
    punctuation = sum(line.count(mark) for mark in "，。！？；：")
    return punctuation > 0 and len(line) >= 4


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.4)
    section.left_margin = section.right_margin = Cm(1.6)


def _add_annotated_line(document: Document, text: str, converter: PinyinConverter, font_size: int = 22) -> int:
    cells = converter.convert(text)
    table = document.add_table(rows=1, cols=max(1, len(cells)))
    table.autofit = True
    for index, cell in enumerate(cells):
        target = table.rows[0].cells[index]
        pinyin_paragraph = target.paragraphs[0]
        pinyin_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pinyin_run = pinyin_paragraph.add_run(cell.pinyin)
        pinyin_run.font.name = "Arial"
        pinyin_run.font.size = Pt(max(8, font_size // 2))
        text_paragraph = target.add_paragraph()
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        text_run = text_paragraph.add_run(cell.text)
        text_run.font.name = "Noto Sans CJK SC"
        text_run.font.size = Pt(font_size)
    return sum(1 for cell in cells if cell.pinyin)


def render_annotated_paragraphs(
    paragraphs: list[str], target_path: Path, converter: PinyinConverter, *, overwrite: bool = False
) -> dict[str, object]:
    """把普通段落生成逐字注音 DOCX。"""

    if target_path.exists() and not overwrite:
        raise FileExistsError(f"目标已存在：{target_path}")
    target_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document)
    hanzi_count = 0
    for paragraph in paragraphs:
        if paragraph.strip():
            hanzi_count += _add_annotated_line(document, paragraph, converter)
    document.save(target_path)
    return {"paragraph_count": len([item for item in paragraphs if item.strip()]), "hanzi_count": hanzi_count, "target": str(target_path)}


def render_poetry(
    poems: list[Poem], target_path: Path, converter: PinyinConverter, *, overwrite: bool = False
) -> dict[str, object]:
    """每首诗独立成页，标题、署名和正文全部使用逐字注音结构。"""

    if target_path.exists() and not overwrite:
        raise FileExistsError(f"目标已存在：{target_path}")
    target_path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    _configure_document(document)
    hanzi_count = 0
    for poem_index, poem in enumerate(poems):
        if poem_index:
            document.add_page_break()
        hanzi_count += _add_annotated_line(document, poem.title, converter, 28)
        hanzi_count += _add_annotated_line(document, poem.attribution, converter, 16)
        for line in poem.lines:
            hanzi_count += _add_annotated_line(document, line, converter, 22)
    document.save(target_path)
    return {"poem_count": len(poems), "hanzi_count": hanzi_count, "target": str(target_path)}


def generate(
    source_path: Path,
    target_path: Path,
    dictionary_path: Path | None = None,
    *,
    overwrite: bool = False,
    poetry: bool = False,
) -> dict[str, object]:
    """执行源文件保护后的普通或诗词注音生成。"""

    source = source_path.resolve()
    target = target_path.resolve()
    if source == target:
        raise ValueError("目标文件不能与源文件相同。")
    converter = PinyinConverter(load_overrides(dictionary_path))
    if poetry:
        result = render_poetry(parse_poetry(source), target, converter, overwrite=overwrite)
    else:
        result = render_annotated_paragraphs(_document_lines(source), target, converter, overwrite=overwrite)
    result["source_sha256"] = hashlib.sha256(source.read_bytes()).hexdigest()
    return result


def main() -> int:
    """提供普通文档、古诗和文言文三种生成模式。"""

    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("target", type=Path)
    parser.add_argument("--dictionary", type=Path)
    parser.add_argument("--mode", choices=("paragraphs", "poetry", "classical"), default="paragraphs")
    parser.add_argument("--overwrite", action="store_true")
    arguments = parser.parse_args()
    if arguments.mode == "classical":
        converter = PinyinConverter(load_overrides(arguments.dictionary))
        articles = parse_classical_chinese(arguments.source)
        poems = [Poem(item.title, item.dynasty, item.author, tuple(item.display_lines(16))) for item in articles]
        result = render_poetry(poems, arguments.target, converter, overwrite=arguments.overwrite)
    else:
        result = generate(
            arguments.source,
            arguments.target,
            arguments.dictionary,
            overwrite=arguments.overwrite,
            poetry=arguments.mode == "poetry",
        )
    print(json.dumps({"status": "completed", **result}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
